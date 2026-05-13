"""Generates Manual_GUI.docx describing the Smart Parking Lot GUI implementation."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x08, 0x91, 0xB2)
MUTED  = RGBColor(0x60, 0x60, 0x60)
DARK   = RGBColor(0x20, 0x20, 0x20)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT
    r.font.name = 'Segoe UI'


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED
    r.font.name = 'Segoe UI'


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK
    r.font.name = 'Segoe UI Semibold'


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    r.font.name = 'Segoe UI Semibold'


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    r.font.name = 'Segoe UI Semibold'


def add_para(doc, text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Segoe UI'
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.runs[0] if p.runs else p.add_run('')
    if not p.runs:
        r = p.add_run(text)
    else:
        p.runs[0].text = text
    r.font.size = Pt(11)
    r.font.name = 'Segoe UI'


def add_bullet_rich(doc, label, body):
    p = doc.add_paragraph(style='List Bullet')
    rl = p.add_run(label)
    rl.bold = True
    rl.font.size = Pt(11)
    rl.font.name = 'Segoe UI'
    rb = p.add_run(body)
    rb.font.size = Pt(11)
    rb.font.name = 'Segoe UI'


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Cascadia Code'
    r.font.size = Pt(9.5)
    r.font.color.rgb = DARK
    # background shading on the run by paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F3F3')
    pPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = 'Segoe UI'
        set_cell_bg(hdr[i], '0891B2')
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.name = 'Segoe UI'


def add_divider(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Build document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Cover
add_title(doc, 'Smart Parking Lot')
add_subtitle(doc, 'Manual de implementación — Interfaz gráfica WinUI 3 nativa Windows 11')
add_para(doc, '')
add_para(doc, 'Versión 1.0 · .NET 10 · Windows App SDK 1.6 · Mayo 2026', italic=True)
add_divider(doc)

# 1. Resumen ejecutivo
add_h1(doc, '1. Resumen ejecutivo')
add_para(doc,
    'Se desarrolló una interfaz gráfica nativa para el sistema Smart Parking Lot, '
    'replicando el 100% de las funciones del menú de consola existente. La GUI fue '
    'construida sobre WinUI 3 con el Windows App SDK 1.6, lo cual garantiza una '
    'apariencia y comportamiento idénticos a las aplicaciones nativas de Windows 11 '
    '(Mica, modo oscuro automático, color de acento del sistema, controles Fluent).')
add_para(doc,
    'La nueva aplicación coexiste con la CLI original: ambas comparten las capas '
    'Core, Application, Hardware y Persistence sin modificaciones. Solo se agregó un '
    'nuevo proyecto SmartParkingLot.Gui que reusa toda la lógica de dominio.')

# 2. Decisiones de diseño
add_h1(doc, '2. Decisiones de diseño')
add_h2(doc, '2.1 Elección del framework')
add_para(doc,
    'El mockup hi-fi (Smart Parking Hi-Fi.html) utiliza el sistema de diseño WinUI 3 / '
    'Fluent Design con Mica, acrílico, segmented controls y la paleta de Windows 11. '
    'Para llevarlo fielmente a una aplicación de escritorio se evaluaron tres opciones:')

add_table(doc,
    ['Framework', 'Fidelidad visual', 'Veredicto'],
    [
      ['WinForms', 'Baja — sin Mica, requiere custom paint', 'Descartado'],
      ['WPF + WPF-UI', 'Alta — aproxima Fluent con librería externa', 'Aceptable'],
      ['WinUI 3 (Windows App SDK)', 'Completa — Mica/acrílico nativos', 'ELEGIDO'],
    ])

add_para(doc,
    'Se eligió WinUI 3 porque ofrece el efecto Mica real, el sistema de temas claro/oscuro '
    'automático del SO, el color de acento del usuario y los controles Fluent (NavigationView, '
    'AutoSuggestBox, ToggleSwitch) sin necesidad de paquetes adicionales.')

add_h2(doc, '2.2 Arquitectura — coexistencia con la CLI')
add_para(doc,
    'En lugar de reemplazar el ParkingLotApp.RunAsync() de la CLI, se extrajo todo el '
    'cableado de dependencias a una clase ParkingServices.BootstrapAsync() que devuelve '
    'objetos listos para usar (ParkingLot, GateController, IParkingRepository, IEventPublisher, '
    'IArduinoReader, etc.). La CLI sigue intacta; la GUI consume los mismos servicios.')
add_para(doc, 'Esto produce tres beneficios:')
add_bullet_rich(doc, 'Reutilización total: ',
    'cero duplicación de lógica de negocio. Cualquier corrección en Core/Application impacta a ambas interfaces.')
add_bullet_rich(doc, 'Pruebas paralelas: ',
    'se puede correr la CLI y la GUI en paralelo apuntando a la misma BD SQLite para verificar la coherencia.')
add_bullet_rich(doc, 'Tolerancia a Arduino ausente: ',
    'si el puerto COM no existe, la GUI muestra "desconectado" en el status bar y sigue funcionando offline; la CLI tenía el mismo comportamiento.')

# 3. Estructura del nuevo proyecto
add_h1(doc, '3. Estructura del nuevo proyecto')
add_para(doc,
    'Se creó la carpeta src/Gui/ con los siguientes archivos clave (todos compilados como '
    'parte del proyecto SmartParkingLot.Gui.csproj):')

add_table(doc,
    ['Archivo', 'Responsabilidad'],
    [
      ['SmartParkingLot.Gui.csproj', 'Proyecto WinUI 3 unpackaged, self-contained, win-x64'],
      ['app.manifest', 'DPI awareness + supported OS para Windows 11'],
      ['App.xaml / App.xaml.cs', 'Bootstrap de la app, splash, captura de excepciones'],
      ['MainWindow.xaml / .cs', 'Shell con title bar, NavigationView, status bar y Mica'],
      ['Styles/Theme.xaml', 'Paleta semántica (Success, Warning, Danger, Tx1–Tx3)'],
      ['Bootstrap/ParkingServices.cs', 'Cableado headless del dominio (mirror de ParkingLotApp)'],
      ['Bootstrap/HardwareConfig.cs', 'Lectura de hardware.json'],
      ['Bootstrap/GuiLogger.cs', 'Logger en memoria con ring-buffer para el log en vivo'],
      ['Bootstrap/GuiConstants.cs', 'Constantes (DEFAULT_LOT_ID, gate IDs, pines...)'],
      ['Controls/Badge.cs', 'Factory de chips Fluent (Success/Warning/Danger/Accent/Neutral)'],
      ['Pages/DashboardPage.*', 'Tiles, ocupación por zona, alertas, resumen gates'],
      ['Pages/MapPage.*', 'Mapa interactivo de spots + control de gates AUTO/MANUAL'],
      ['Pages/LogPage.*', 'Búsqueda en historial (placas, sensores, dispositivos)'],
      ['Pages/AdminPage.*', 'Lista completa de spots desde la BD'],
      ['Pages/HardwarePage.*', 'Conexión Arduino, log en vivo, simulación de sensores'],
      ['Pages/SplashPage.cs', 'Splash + página de Configuración'],
      ['hardware.json', 'Mapeo de sensores → spots y gates → pines'],
    ])

# 4. Cobertura funcional
add_h1(doc, '4. Cobertura funcional — CLI ↔ GUI')
add_para(doc,
    'Las 11 funciones del menú de consola están todas accesibles desde la GUI. La tabla '
    'siguiente muestra el mapeo función → ubicación en la interfaz gráfica.')

add_table(doc,
    ['Función del menú CLI', 'Ubicación en la GUI'],
    [
      ['Solicitar entrada de vehículo', 'Mapa de Spots → botón "Solicitar entrada"'],
      ['Solicitar salida de vehículo', 'Mapa de Spots → botón "Solicitar salida"'],
      ['Actualizar estado de espacio (sensor manual)', 'Mapa de Spots → click directo sobre el spot'],
      ['Ver estado del parqueadero', 'Dashboard (tiles + ocupación por zona)'],
      ['Ver historial de un vehículo', 'Historial → tipo "Solicitudes (placa)"'],
      ['Ver lecturas de un sensor', 'Historial → tipo "Lecturas de sensor"'],
      ['Ver acciones de un dispositivo', 'Historial → tipo "Acciones de dispositivo"'],
      ['Monitoreo en tiempo real (Arduino)', 'Hardware → panel "Log en vivo"'],
      ['Ver estado de espacios (BD)', 'Gestión de Spots (tabla completa con filtros)'],
      ['Ver logs recientes', 'Hardware → "Archivo de log de hoy"'],
      ['Simular sensor de puerta (IR)', 'Mapa o Hardware → "Simular IR entrada/salida"'],
    ])

# 5. Páginas de la aplicación
add_h1(doc, '5. Recorrido por las páginas')

add_h2(doc, '5.1 Dashboard')
add_para(doc, 'Pantalla principal con cuatro indicadores:')
add_bullet(doc, 'Spots ocupados / Disponibles / Solicitudes hoy / % Ocupación')
add_bullet(doc, 'Barras de progreso por zona con color semántico (verde < 70% < amarillo < 90% < rojo)')
add_bullet(doc, 'Lista de alertas recientes leídas del GuiLogger (entradas Warning/Error)')
add_bullet(doc, 'Tarjetas resumen de cada gate con su estado y tipo (ENTRY/EXIT)')

add_h2(doc, '5.2 Mapa de Spots')
add_para(doc,
    'Núcleo interactivo de la aplicación. Renderiza los spots agrupados por zona usando '
    'VariableSizedWrapGrid. Cada celda es clicable: al hacer click se publica un '
    'SensorReadingReceived en el bus de eventos, exactamente igual que en la CLI.')
add_para(doc, 'En el panel derecho se controlan los gates:')
add_bullet(doc, 'Cambio de modo AUTOMATIC ↔ MANUAL por gate')
add_bullet(doc, 'Botón Abrir/Cerrar manual cuando está en modo MANUAL')
add_bullet(doc, 'Badge con el estado actual (Abierta/Cerrada)')
add_para(doc,
    'Los dos botones de la command bar (Solicitar entrada / Solicitar salida) abren un '
    'ContentDialog que pide la placa, crean un EntryRequest o ExitRequest y lo despachan '
    'al GateController. La respuesta (concedido/denegado) se muestra en otro ContentDialog.')

add_h2(doc, '5.3 Historial')
add_para(doc,
    'Consulta unificada a IParkingRepository con tres modos según el ComboBox:')
add_bullet(doc, 'Solicitudes por placa: GetRequestHistoryAsync')
add_bullet(doc, 'Lecturas de sensor: GetSensorReadingsAsync')
add_bullet(doc, 'Acciones de dispositivo: GetDeviceActionsAsync')
add_para(doc,
    'Los resultados se renderizan en filas con timestamp, badge de tipo, detalle y '
    'referencia (RequestId / SensorId / ActionId).')

add_h2(doc, '5.4 Gestión de Spots')
add_para(doc,
    'Lista todos los spots de la base de datos (no del lote en memoria) usando '
    'GetSpotsByLotIdAsync. Soporta búsqueda por texto sobre ID/tipo/dirección y filtro '
    'por tipo (Estándar/PMR/Moto). Footer con el conteo "X de Y spots".')

add_h2(doc, '5.5 Hardware / Arduino')
add_para(doc, 'Centro de control y diagnóstico del bridge serial:')
add_bullet(doc, 'Estado de conexión (Conectado/Desconectado) con badge animado')
add_bullet(doc, 'Botón Conectar/Desconectar que llama Bridge.StartListening / StopListening')
add_bullet(doc, 'Escaneo de puertos COM disponibles vía SerialPort.GetPortNames()')
add_bullet(doc, 'Log en vivo suscrito al GuiLogger, con toggle de auto-scroll')
add_bullet(doc, 'Publicación manual de eventos de sensor (cualquier sensor del sistema)')
add_bullet(doc, 'Simulación de IR para entrada/salida (mismo evento que la CLI)')
add_bullet(doc, 'Tail del archivo de log físico del día (últimas 80 líneas)')

# 6. Detalles técnicos
add_h1(doc, '6. Detalles técnicos relevantes')

add_h2(doc, '6.1 Mica y title bar custom')
add_para(doc,
    'La ventana usa Mica como backdrop nativo cuando el SO lo soporta '
    '(Windows 11), con fallback a acrílico. El title bar es un Grid custom '
    'declarado en XAML y registrado con ExtendsContentIntoTitleBar = true + '
    'SetTitleBar(AppTitleBar), lo que produce el área de arrastre nativa sin '
    'redibujar nada.')
add_code(doc,
    'private void TryApplyBackdrop()\n'
    '{\n'
    '    if (MicaController.IsSupported())\n'
    '        SystemBackdrop = new MicaBackdrop { Kind = MicaKind.Base };\n'
    '    else if (DesktopAcrylicController.IsSupported())\n'
    '        SystemBackdrop = new DesktopAcrylicBackdrop();\n'
    '}')

add_h2(doc, '6.2 Paleta semántica con ThemeDictionaries')
add_para(doc,
    'En Styles/Theme.xaml se definen dos diccionarios — Light y Dark — con la misma lista '
    'de claves (SuccessBrush, WarningBrush, DangerBrush, Tx1Brush a Tx3Brush, Layer1Brush, '
    'Layer2Brush, etc.). WinUI cambia automáticamente el set activo según el tema del SO.')

add_h2(doc, '6.3 Self-contained deployment')
add_para(doc,
    'El .csproj usa WindowsAppSDKSelfContained=true + SelfContained=true para que el '
    'ejecutable final incluya tanto el runtime de .NET 10 como el de Windows App SDK 1.6. '
    'Esto evita el error clásico "This application requires the Windows App Runtime 1.6". '
    'El binario pesa ~150 MB pero corre en cualquier Windows 10/11 x64 sin instalar nada.')

add_h2(doc, '6.4 Captura de excepciones tempranas')
add_para(doc,
    'App.OnLaunched envuelve la construcción de MainWindow y el bootstrap en try/catch '
    'separados, y App.UnhandledException tiene un fallback que muestra un MessageBox Win32 '
    'cuando todavía no hay ventana WinUI donde mostrar el error. Cada paso del arranque '
    'se registra en %TEMP%\\SmartParkingLot.Gui-crash.log para diagnóstico.')

# 7. Cómo ejecutar
add_h1(doc, '7. Cómo compilar y ejecutar')

add_h3(doc, 'Opción A — Desde la línea de comandos')
add_code(doc,
    'cd C:\\Dev\\smart-parking-lot\\src\\Gui\n'
    'dotnet run -c Debug -r win-x64')

add_h3(doc, 'Opción B — Ejecutable directo')
add_para(doc, 'Después de compilar una vez, basta doble click sobre:')
add_code(doc,
    'C:\\Dev\\smart-parking-lot\\src\\Gui\\bin\\Debug\\\n'
    'net10.0-windows10.0.19041.0\\win-x64\\SmartParkingLot.Gui.exe')

add_h3(doc, 'Opción C — Visual Studio')
add_para(doc,
    'Abrir smart-parking-lot.sln, marcar SmartParkingLot.Gui como Startup Project, '
    'cambiar plataforma de Any CPU a x64 y pulsar F5.')

# 8. Cierre
add_h1(doc, '8. Notas finales')
add_bullet(doc,
    'La GUI no rompe la CLI: ambos proyectos compilan y se ejecutan independientemente.')
add_bullet(doc,
    'src/Gui/hardware.json tiene un mapeo expandido (9 spots) para que el mapa luzca; '
    'en producción debe alinearse con el cableado real del Arduino.')
add_bullet(doc,
    'Los HTMLs del mockup se conservaron en src/Gui/ como referencia visual; no se '
    'incluyen en la compilación.')
add_bullet(doc,
    'Si la ventana no aparece al ejecutar, revisar %TEMP%\\SmartParkingLot.Gui-crash.log '
    'donde se registra cada paso del arranque y cualquier excepción.')

doc.save(r'C:\Dev\smart-parking-lot\docs\Manual_GUI.docx')
print('OK: Manual_GUI.docx generado')
